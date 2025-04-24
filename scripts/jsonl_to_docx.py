#!/usr/bin/env python3
import argparse
import json
import os
import glob
import logging
import re
from docx import Document
from docx.shared import Pt

# Basic logging setup
log_format = '%(asctime)s - %(name)s - %(levelname)s - %(message)s'
logging.basicConfig(level=logging.INFO, format=log_format)
logger = logging.getLogger(__name__)

def parse_args():
    parser = argparse.ArgumentParser(description="Read local JSONL files, extract text, and write to local .docx files.")
    parser.add_argument(
        "jsonl_dir",
        help="Directory containing the input JSONL files.",
    )
    parser.add_argument(
        "output_dir",
        help="Local directory to store output .docx files.",
    )
    return parser.parse_args()

def sanitize_for_xml(text):
    """
    Removes characters that are invalid in XML 1.0, except for tab, newline, and carriage return.
    Also removes null bytes.
    """
    if not isinstance(text, str):
        return text # Return as is if not a string

    # Remove null bytes explicitly
    text = text.replace('\x00', '')

    # Regex to match invalid XML characters (control characters excluding \t, \n, \r)
    invalid_xml_chars_re = re.compile(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F-\x84\x86-\x9F]')
    return invalid_xml_chars_re.sub('', text)

def main():
    args = parse_args()

    # Ensure local output directory exists
    os.makedirs(args.output_dir, exist_ok=True)
    logger.info(f"Output directory set to: {args.output_dir}")

    # Find all JSONL files in the input directory
    jsonl_files = glob.glob(os.path.join(args.jsonl_dir, "*.jsonl"))
    logger.info(f"Found {len(jsonl_files)} JSONL files in {args.jsonl_dir}")

    if not jsonl_files:
        print("No JSONL files found in the specified directory.")
        return

    for jsonl_path in jsonl_files:
        logger.info(f"Processing JSONL file: {jsonl_path}")
        source_file_base = os.path.splitext(os.path.basename(jsonl_path))[0].replace('_output', '') # Default base name

        try:
            document = Document()
            # Optional: Set default font (if needed)
            # style = document.styles['Normal']
            # font = style.font
            # font.name = 'Arial' # Or another preferred font
            # font.size = Pt(11)

            with open(jsonl_path, 'r', encoding='utf-8') as f_in:
                for i, line in enumerate(f_in):
                    if not line.strip():
                        continue
                    try:
                        record = json.loads(line)
                        text_content = record.get("text", "")
                        metadata = record.get("metadata", {})
                        # Use source file from metadata if available for better naming
                        source_file_meta = metadata.get("Source-File")
                        if source_file_meta:
                             source_file_base = os.path.splitext(os.path.basename(source_file_meta))[0]
                             source_file_base = "".join(c if c.isalnum() or c in ('-', '_') else '_' for c in source_file_base) # Sanitize

                        if text_content:
                            # Add paragraph for each non-empty page text
                            sanitized_content = sanitize_for_xml(text_content)
                            document.add_paragraph(sanitized_content)
                            # Add a page break after each page's content? Optional.
                            # document.add_page_break()

                    except json.JSONDecodeError:
                        logger.warning(f"Failed to decode JSON line {i+1} in {jsonl_path}")
                        continue

            if len(document.paragraphs) > 0: # Only save if content was added
                output_filename = f"{source_file_base}.docx"
                output_path = os.path.join(args.output_dir, output_filename)
                try:
                    document.save(output_path)
                    logger.info(f"Successfully wrote DOCX to {output_path}")
                except Exception as e:
                    logger.error(f"Failed to save DOCX file {output_path}: {e}")
            else:
                 logger.info(f"No text content found in {jsonl_path}, skipping DOCX creation.")


        except Exception as e:
            logger.error(f"Failed to process file {jsonl_path}: {e}")

    logger.info("Done processing all JSONL files.")

if __name__ == "__main__":
    main()